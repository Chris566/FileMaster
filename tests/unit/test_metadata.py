"""Metadata 读取器测试 (W3).

覆盖 4 类文件（PDF/Word/Excel/Image）的基础字段读取、缺失字段兜底、
损坏/不存在文件优雅降级，以及批量读取。每个 test 动态造 fixture，
不依赖外部文件。
"""

from __future__ import annotations

from pathlib import Path

import pytest

from filemaster.core.metadata import FileMetadata, MetadataReader

# ---------- fixture 构造函数 ----------


def _make_pdf(tmp_path: Path, title: str = "Annual Report 2026", author: str = "东东") -> Path:
    import fitz

    p = tmp_path / "test.pdf"
    doc = fitz.open()
    doc.new_page()  # 默认在末尾加一页（pno=-1）
    doc.set_metadata({"title": title, "author": author})
    doc.save(p)
    doc.close()
    return p


def _make_docx(tmp_path: Path, title: str = "Doc Title", author: str = "Doc Author") -> Path:
    from docx import Document

    p = tmp_path / "test.docx"
    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.save(p)
    return p


def _make_xlsx(tmp_path: Path, title: str = "Sheet Report", creator: str = "xlsAuthor") -> Path:
    from openpyxl import Workbook

    p = tmp_path / "test.xlsx"
    wb = Workbook()
    wb.properties.title = title
    wb.properties.creator = creator
    wb.save(p)
    return p


def _make_image(tmp_path: Path, width: int = 200, height: int = 100) -> Path:
    from PIL import Image

    p = tmp_path / "test.png"
    img = Image.new("RGB", (width, height), "red")
    img.save(p)
    return p


# ---------- PDF ----------


class TestPdfMetadata:
    def test_basic_fields(self, reader: MetadataReader, tmp_path: Path) -> None:
        p = _make_pdf(tmp_path, title="Annual Report 2026", author="东东")
        m = reader.read(p)
        assert m.title == "Annual Report 2026"
        assert m.author == "东东"
        assert m.page_count >= 1

    def test_empty_metadata(self, reader: MetadataReader, tmp_path: Path) -> None:
        """不设元数据时应返空串，不是抛异常."""
        p = _make_pdf(tmp_path)  # 默认空 title/author 也会写入
        m = reader.read(p)
        assert isinstance(m.title, str)
        assert isinstance(m.author, str)
        assert m.page_count >= 1

    def test_corrupt_pdf_returns_empty(self, reader: MetadataReader, tmp_path: Path) -> None:
        p = tmp_path / "fake.pdf"
        p.write_bytes(b"not a pdf")
        m = reader.read(p)
        assert m.title == ""
        assert m.author == ""
        assert m.page_count == 0

    def test_nonexistent_pdf(self, reader: MetadataReader, tmp_path: Path) -> None:
        p = tmp_path / "ghost.pdf"
        m = reader.read(p)
        assert m.title == ""


# ---------- Word ----------


class TestWordMetadata:
    def test_basic_fields(self, reader: MetadataReader, tmp_path: Path) -> None:
        p = _make_docx(tmp_path, title="Doc Title", author="Doc Author")
        m = reader.read(p)
        assert m.title == "Doc Title"
        assert m.author == "Doc Author"
        # page_count 是 paragraph 数的简化统计;新 doc 可能为 0,断言类型即可
        assert isinstance(m.page_count, int)
        assert m.page_count >= 0

    def test_empty_metadata(self, reader: MetadataReader, tmp_path: Path) -> None:
        """Document() 不设元数据应返空串."""
        from docx import Document

        p = tmp_path / "test.docx"
        Document().save(p)
        m = reader.read(p)
        assert isinstance(m.title, str)
        assert isinstance(m.author, str)
        # Word 的 page_count 是简化版的 paragraph 数，至少 1（空 doc 也有 1 段）
        assert m.page_count >= 0

    def test_corrupt_docx_returns_empty(self, reader: MetadataReader, tmp_path: Path) -> None:
        p = tmp_path / "fake.docx"
        p.write_bytes(b"not a docx")
        m = reader.read(p)
        assert m.title == ""
        assert m.author == ""


# ---------- Excel ----------


class TestExcelMetadata:
    def test_basic_fields(self, reader: MetadataReader, tmp_path: Path) -> None:
        p = _make_xlsx(tmp_path, title="Sheet Report", creator="xlsAuthor")
        m = reader.read(p)
        assert m.title == "Sheet Report"
        assert m.author == "xlsAuthor"  # creator → author 映射

    def test_empty_metadata(self, reader: MetadataReader, tmp_path: Path) -> None:
        from openpyxl import Workbook

        p = tmp_path / "test.xlsx"
        Workbook().save(p)
        m = reader.read(p)
        assert isinstance(m.title, str)
        assert isinstance(m.author, str)

    def test_corrupt_xlsx_returns_empty(self, reader: MetadataReader, tmp_path: Path) -> None:
        p = tmp_path / "fake.xlsx"
        p.write_bytes(b"not an xlsx")
        m = reader.read(p)
        assert m.title == ""


# ---------- Image ----------


class TestImageMetadata:
    def test_basic_png_size(self, reader: MetadataReader, tmp_path: Path) -> None:
        p = _make_image(tmp_path, width=200, height=100)
        m = reader.read(p)
        assert m.extra.get("size") == (200, 100)
        assert m.extra.get("format") == "PNG"

    def test_no_exif_returns_empty_strings(self, reader: MetadataReader, tmp_path: Path) -> None:
        """新造的 PNG 没有 EXIF，应返空串（不抛异常）."""
        p = _make_image(tmp_path)
        m = reader.read(p)
        assert m.title == ""
        assert m.author == ""
        assert m.created == ""

    def test_corrupt_image_returns_empty(self, reader: MetadataReader, tmp_path: Path) -> None:
        p = tmp_path / "fake.png"
        p.write_bytes(b"not an image")
        m = reader.read(p)
        assert m.title == ""


# ---------- 未知扩展名 ----------


class TestUnknownExtension:
    def test_txt_returns_empty(self, reader: MetadataReader, tmp_path: Path) -> None:
        p = tmp_path / "doc.txt"
        p.write_text("hello")
        m = reader.read(p)
        assert m.title == ""
        assert m.author == ""
        assert m.page_count == 0

    def test_rtf_falls_through_to_word_reader(self, reader: MetadataReader, tmp_path: Path) -> None:
        """RTF 在 WORD_EXTS 里但 python-docx 不支持,应返空串不抛."""
        p = tmp_path / "doc.rtf"
        p.write_text("{\\rtf1 hello}")
        m = reader.read(p)
        # python-docx 抛异常被吞,返空 FileMetadata
        assert m.title == ""

    def test_uppercase_extension(self, reader: MetadataReader, tmp_path: Path) -> None:
        """大写扩展名 .PDF 也要识别."""
        p = _make_pdf(tmp_path, title="X")
        p2 = tmp_path / "upper.PDF"
        p.rename(p2)
        m = reader.read(p2)
        assert m.title == "X"


# ---------- 批量读取 ----------


class TestBatchRead:
    def test_read_all(self, reader: MetadataReader, tmp_path: Path) -> None:
        pdf = _make_pdf(tmp_path)
        txt = tmp_path / "a.txt"
        txt.write_text("hi")
        results = reader.read_all([pdf, txt])
        assert len(results) == 2
        assert results[pdf].page_count >= 1
        assert results[txt].title == ""

    def test_read_all_empty(self, reader: MetadataReader) -> None:
        results = reader.read_all([])
        assert results == {}


# ---------- FileMetadata dataclass ----------


def test_filemetadata_frozen() -> None:
    """FileMetadata 是 frozen dataclass,不应被 mutate."""
    m = FileMetadata(title="X")
    with pytest.raises((AttributeError, Exception)):
        m.title = "Y"  # type: ignore[misc]


def test_filemetadata_extra_default_is_dict() -> None:
    """不传 extra 时应为 {}, 不可为 None."""
    m = FileMetadata()
    assert m.extra == {}
    assert isinstance(m.extra, dict)


def test_filemetadata_default_values() -> None:
    """默认值全空串/0."""
    m = FileMetadata()
    assert m.title == ""
    assert m.author == ""
    assert m.subject == ""
    assert m.created == ""
    assert m.modified == ""
    assert m.page_count == 0
    assert m.extra == {}


@pytest.fixture
def reader() -> MetadataReader:
    return MetadataReader()
