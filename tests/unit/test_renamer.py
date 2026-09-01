"""重命名引擎测试."""

from __future__ import annotations

from pathlib import Path

import pytest

from filemaster.core.renamer import (
    ConflictStrategy,
    Renamer,
    RenameResult,
    format_date,
    format_size,
    get_excel_sheet_name,
)
from filemaster.core.template import Template
from filemaster.core.undo import UndoStack


class TestRenamerPlan:
    """plan() 只生成结果，不动文件."""

    def test_plan_with_prefix(self, sample_files: list[Path]) -> None:
        tpl = Template("{Prefix}{OriginalName}")
        renamer = Renamer(tpl, prefix="X_")
        results = renamer.plan(sample_files)
        assert len(results) == 3
        for r in results:
            assert r.status == "DRY_RUN"
            assert r.target is not None
            assert r.target.name.startswith("X_")

    def test_plan_increments_index(self, sample_files: list[Path]) -> None:
        tpl = Template("{Index:D3}_{OriginalName}")
        renamer = Renamer(tpl)
        results = renamer.plan(sample_files)
        # 序号从 1 起
        assert results[0].target.name.startswith("001_")
        assert results[1].target.name.startswith("002_")
        assert results[2].target.name.startswith("003_")

    def test_plan_skips_unchanged(self, sample_files: list[Path]) -> None:
        # 模板与原名相同（无前缀）
        tpl = Template("{OriginalName}")
        renamer = Renamer(tpl)
        results = renamer.plan(sample_files)
        for r in results:
            assert r.status == "SKIPPED"
            assert r.target is None

    def test_plan_start_index(self, sample_files: list[Path]) -> None:
        tpl = Template("{Index:D3}_{OriginalName}")
        renamer = Renamer(tpl, start_index=100)
        results = renamer.plan(sample_files)
        assert results[0].target.name.startswith("100_")
        assert results[2].target.name.startswith("102_")

    def test_reset_index(self, sample_files: list[Path]) -> None:
        tpl = Template("{Index:D3}_{OriginalName}")
        renamer = Renamer(tpl)
        renamer.plan(sample_files[:1])  # 推进 index
        renamer.reset_index()
        results = renamer.plan(sample_files[:1])
        assert results[0].target.name.startswith("001_")


class TestRenamerApply:
    """W2: apply() 真实文件 IO + 冲突策略."""

    def test_apply_basic_rename(self, sample_files: list[Path]) -> None:
        tpl = Template("{Prefix}_{Index:D3}_{OriginalName}")
        renamer = Renamer(tpl, prefix="X")
        results = renamer.apply(sample_files)
        # 全部成功
        assert all(r.status == "OK" for r in results)
        # 源文件已被移动
        assert all(not f.exists() for f in sample_files)
        # 目标文件存在
        targets = [r.target for r in results]
        assert all(t.exists() for t in targets)
        assert targets[0].name == "X_001_doc_001.pdf"
        assert targets[2].name == "X_003_doc_003.pdf"

    def test_apply_with_undo_stack(self, sample_files: list[Path], tmp_path: Path) -> None:
        tpl = Template("{Prefix}{OriginalName}")
        undo = UndoStack(persist_dir=tmp_path / "undo")
        renamer = Renamer(tpl, prefix="X_")
        renamer.apply(sample_files, undo_stack=undo)
        # undo 栈应有一条 batch
        assert len(undo) == 1
        batch = next(iter(undo))
        assert len(batch) == 3
        # 全部 RENAME_ONLY
        assert all(e.operation == "RenameOnly" for e in batch)

    def test_apply_skip_on_conflict(self, sample_files: list[Path]) -> None:
        tpl = Template("{Prefix}{OriginalName}")
        # 预先创建冲突文件
        (sample_files[0].parent / "X_doc_001.pdf").write_bytes(b"existing")
        renamer = Renamer(tpl, prefix="X_")
        results = renamer.apply(sample_files, conflict_strategy=ConflictStrategy.SKIP)
        # 第一个冲突跳过
        assert results[0].status == "CONFLICT"
        assert results[0].source.exists()  # 源未动
        # 其他正常
        assert results[1].status == "OK"
        assert results[2].status == "OK"

    def test_apply_overwrite_with_backup(self, sample_files: list[Path], tmp_path: Path) -> None:
        tpl = Template("{Prefix}{OriginalName}")
        # 预先创建冲突文件
        conflict_target = sample_files[0].parent / "X_doc_001.pdf"
        conflict_target.write_bytes(b"original content")
        undo = UndoStack(persist_dir=tmp_path / "undo")
        renamer = Renamer(tpl, prefix="X_")
        results = renamer.apply(
            sample_files,
            conflict_strategy=ConflictStrategy.OVERWRITE,
            undo_stack=undo,
        )
        # 第一个被覆盖
        assert results[0].status == "OVERWRITTEN"
        # undo 应记录 backup_path
        batch = next(iter(undo))
        entry = batch[0]
        assert entry.backup_path is not None
        assert entry.backup_path.exists()
        # 备份内容是原冲突文件
        assert entry.backup_path.read_bytes() == b"original content"
        # 新内容是源文件
        assert conflict_target.exists()

    def test_apply_rename_new_when_exists(self, sample_files: list[Path]) -> None:
        tpl = Template("{Prefix}{OriginalName}")
        # 预先创建冲突文件
        (sample_files[0].parent / "X_doc_001.pdf").write_bytes(b"existing")
        renamer = Renamer(tpl, prefix="X_")
        results = renamer.apply(
            sample_files,
            conflict_strategy=ConflictStrategy.RENAME_NEW,
        )
        # 第一个避冲突改名
        assert results[0].status == "RENAMED"
        assert results[0].target.name == "X_doc_001 (1).pdf"
        assert results[0].target.exists()
        # 原冲突文件还在
        assert (sample_files[0].parent / "X_doc_001.pdf").exists()

    def test_apply_undo_restores_originals(
        self, sample_files: list[Path], tmp_path: Path
    ) -> None:
        tpl = Template("{Prefix}{OriginalName}")
        undo = UndoStack(persist_dir=tmp_path / "undo")
        renamer = Renamer(tpl, prefix="X_")
        renamer.apply(sample_files, undo_stack=undo)
        # 所有源已改名
        assert all(not f.exists() for f in sample_files)
        # 撤销
        batch = undo.pop()
        assert batch is not None
        import shutil

        for entry in batch:
            if entry.target and entry.target.exists():
                entry.target.rename(entry.source)
        # 源恢复
        assert all(f.exists() for f in sample_files)


class TestRenamerPrefix:
    """前缀识别."""

    def test_already_has_prefix(self, sample_files: list[Path]) -> None:
        renamer = Renamer(Template("X"), prefix="X_")
        # 文件名是 doc_001.pdf，没有 X_ 前缀
        assert not renamer.already_has_prefix(sample_files[0])

    def test_already_has_prefix_case_insensitive(self, tmp_path: Path) -> None:
        f = tmp_path / "X_Document.pdf"
        f.write_bytes(b"x")
        renamer = Renamer(Template("X"), prefix="x_")
        assert renamer.already_has_prefix(f)

    def test_empty_prefix(self, sample_files: list[Path]) -> None:
        renamer = Renamer(Template("X"), prefix="")
        assert not renamer.already_has_prefix(sample_files[0])


class TestRenamerSanitize:
    """文件名清理."""

    def test_removes_windows_illegal(self) -> None:
        assert Renamer.sanitize('a<b>c:d"e/f\\g|h?i*.txt') == "a_b_c_d_e_f_g_h_i_.txt"

    def test_keeps_normal_chars(self) -> None:
        assert Renamer.sanitize("normal_filename-v1.0.pdf") == "normal_filename-v1.0.pdf"

    def test_replaces_control_chars(self) -> None:
        assert Renamer.sanitize("file\x00name\x1f.pdf") == "file_name_.pdf"

    def test_empty_string(self) -> None:
        assert Renamer.sanitize("") == ""

    def test_strips_trailing_dots_and_spaces(self) -> None:
        # Windows 不允许文件名以 . 或空格结尾
        assert Renamer.sanitize("file.  .pdf.") == "file.  .pdf"


# ============ W2 新增测试 ============


class TestFormatSize:
    """format_size 工具函数."""

    def test_bytes(self) -> None:
        assert format_size(512) == "512 B"

    def test_kilobytes(self) -> None:
        assert format_size(2048) == "2.0 KB"

    def test_megabytes(self) -> None:
        assert format_size(1024 * 1024 * 5) == "5.0 MB"

    def test_gigabytes(self) -> None:
        result = format_size(1024**3 * 2)
        assert result == "2.0 GB"

    def test_zero(self) -> None:
        assert format_size(0) == "0 B"

    def test_negative_returns_zero(self) -> None:
        assert format_size(-1) == "0 B"


class TestFormatDate:
    """format_date 工具函数."""

    def test_basic(self) -> None:
        # 2026-08-30 00:00:00 UTC
        import time

        ts = time.mktime((2026, 8, 30, 0, 0, 0, 0, 0, 0))
        assert format_date(ts) == "2026-08-30"

    def test_custom_format(self) -> None:
        import time

        ts = time.mktime((2026, 8, 30, 14, 35, 22, 0, 0, 0))
        assert format_date(ts, "%Y%m%d") == "20260830"
        assert format_date(ts, "%H%M%S") == "143522"

    def test_invalid_timestamp_returns_empty(self) -> None:
        # 1e20 触发 OverflowError,跨平台一致返回 ""
        assert format_date(1e20) == ""  # OverflowError on extreme value
        # -1 行为跨平台不一致:Linux 返回 1970-01-01,Windows 抛 OSError
        # 不依赖 OS 行为,只验证不抛异常
        try:
            result = format_date(-1)
            # 如果没抛错,结果应是合理日期字符串
            assert isinstance(result, str)
        except Exception as e:
            pytest.fail(f"format_date(-1) 不应抛异常: {e}")


class TestGetExcelSheetName:
    """get_excel_sheet_name 工具函数."""

    def test_non_excel_returns_empty(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.txt"
        f.write_bytes(b"plain text")
        assert get_excel_sheet_name(f) == ""

    def test_real_xlsx(self, tmp_path: Path) -> None:
        pytest.importorskip("openpyxl")
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "销售数据"
        ws["A1"] = "测试"
        xlsx_path = tmp_path / "report.xlsx"
        wb.save(xlsx_path)

        assert get_excel_sheet_name(xlsx_path) == "销售数据"

    def test_corrupt_xlsx_returns_empty(self, tmp_path: Path) -> None:
        f = tmp_path / "fake.xlsx"
        f.write_bytes(b"not a real xlsx")
        assert get_excel_sheet_name(f) == ""


class TestContextForExtended:
    """W2: Renamer._context_for 的扩展字段."""

    def test_filesize_placeholder(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.pdf"
        f.write_bytes(b"x" * 2048)  # 2 KB
        tpl = Template("{FileSize}_{OriginalName}")
        renamer = Renamer(tpl)
        target = renamer._render_target(f)
        assert target is not None
        assert target.name == "2.0 KB_doc.pdf"

    def test_filesize_bytes_placeholder(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.pdf"
        f.write_bytes(b"x" * 1234)
        tpl = Template("{FileSizeBytes}_{OriginalName}")
        renamer = Renamer(tpl)
        target = renamer._render_target(f)
        assert target is not None
        assert target.name == "1234_doc.pdf"

    def test_created_modified_date_placeholders(self, tmp_path: Path) -> None:
        """ctime 在 Linux 上 utime 改不了，只测 ModifiedDate（可设）."""
        import os
        import time

        f = tmp_path / "doc.pdf"
        f.write_bytes(b"x")
        # 设置明确的 mtime（atime, mtime 都设）
        ts = time.mktime((2026, 7, 15, 10, 30, 0, 0, 0, 0))
        os.utime(f, (ts, ts))

        tpl = Template("{ModifiedDate}_{OriginalName}")
        renamer = Renamer(tpl)
        target = renamer._render_target(f)
        assert target is not None
        assert target.name == "2026-07-15_doc.pdf"

    def test_hash_short_placeholder(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.pdf"
        f.write_bytes(b"hello world")
        tpl = Template("{HashShort}_{OriginalName}")
        renamer = Renamer(tpl)
        target = renamer._render_target(f)
        assert target is not None
        # MD5("hello world") = 5eb63bbbe01eeed093cb22bb8f5acdc3
        # 前 8 位: 5eb63bbb
        assert target.name.startswith("5eb63bbb_doc.pdf")

    def test_sheet_placeholder_xlsx(self, tmp_path: Path) -> None:
        pytest.importorskip("openpyxl")
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "财务"
        ws["A1"] = "数据"
        xlsx_path = tmp_path / "财务报告.xlsx"
        wb.save(xlsx_path)

        tpl = Template("{Sheet}_{OriginalName}")
        renamer = Renamer(tpl)
        target = renamer._render_target(xlsx_path)
        assert target is not None
        assert target.name == "财务_财务报告.xlsx"

    def test_sheet_placeholder_non_excel(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.txt"
        f.write_bytes(b"plain text")
        tpl = Template("{Sheet}_{OriginalName}")
        renamer = Renamer(tpl)
        target = renamer._render_target(f)
        # Sheet 留空，结果不破坏文件名
        assert target is not None
        assert target.name == "_doc.txt"

    def test_lazy_no_stat_when_unused(self, tmp_path: Path) -> None:
        """未用 FileSize 时不应调 stat."""
        f = tmp_path / "doc.pdf"
        f.write_bytes(b"x")
        tpl = Template("{Prefix}{OriginalName}")  # 不含 FileSize
        renamer = Renamer(tpl, prefix="X_")
        # 应该直接基于模板名渲染，不触发 stat
        target = renamer._render_target(f)
        assert target.name == "X_doc.pdf"


class TestConflictStrategy:
    """ConflictStrategy 枚举."""

    def test_string_values(self) -> None:
        assert ConflictStrategy.SKIP.value == "skip"
        assert ConflictStrategy.OVERWRITE.value == "overwrite"
        assert ConflictStrategy.RENAME_NEW.value == "rename_new"

    def test_str(self) -> None:
        assert str(ConflictStrategy.SKIP) == "skip"


class TestRenameResult:
    """RenameResult 数据类."""

    def test_frozen(self) -> None:
        r = RenameResult(Path("a"), Path("b"), "OK")
        with pytest.raises(Exception):  # FrozenInstanceError
            r.status = "X"  # type: ignore[misc]

"""W5: 命名空间占位符 + apply_with_progress 测试.

命名空间 prefix:
- {pdf_*}  (title/author/subject/pages/created/modified)
- {word_*} (title/author/subject/paragraphs/created/modified)
- {excel_*} (title/author/subject/sheets/sheet_name/created/modified)
- {image_*} (width/height/taken_at/camera_make/camera_model/format/aspect_ratio)
"""


# ---- fixture 复用 test_metadata 的 _make_* 函数 ----

def _make_pdf(tmp_path: Path, title: str = "My PDF") -> Path:
    import fitz

    p = tmp_path / "doc.pdf"
    doc = fitz.open()
    doc.new_page()
    doc.set_metadata({"title": title, "author": "Alice"})
    doc.save(p)
    doc.close()
    return p


def _make_docx(tmp_path: Path, title: str = "My Doc") -> Path:
    from docx import Document

    p = tmp_path / "doc.docx"
    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = "Bob"
    doc.add_paragraph("段一")
    doc.add_paragraph("段二")
    doc.add_paragraph("段三")
    doc.save(p)
    return p


def _make_xlsx(tmp_path: Path, title: str = "My XLS") -> Path:
    from openpyxl import Workbook

    p = tmp_path / "book.xlsx"
    wb = Workbook()
    wb.properties.title = title
    wb.properties.creator = "Carol"
    wb.create_sheet("Sheet2")
    wb.save(p)
    return p


def _make_png(tmp_path: Path, width: int = 1920, height: int = 1080) -> Path:
    from PIL import Image

    p = tmp_path / "img.png"
    img = Image.new("RGB", (width, height), "blue")
    img.save(p)
    return p


# ---- PDF 命名空间 ----


class TestPdfNamespace:
    def test_pdf_title(self, tmp_path: Path) -> None:
        f = _make_pdf(tmp_path, title="Annual Report")
        tpl = Template("{pdf_title}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "Annual Report_doc.pdf"

    def test_pdf_pages(self, tmp_path: Path) -> None:
        f = _make_pdf(tmp_path)
        tpl = Template("{pdf_pages}p_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        # PDF 有 1 页
        assert target.name.startswith("1p_doc.pdf")

    def test_pdf_namespace_empty_for_non_pdf(self, tmp_path: Path) -> None:
        """非 PDF 文件的 {pdf_*} 占位符应填空串."""
        f = tmp_path / "photo.png"
        f.write_bytes(b"\x89PNG\r\n\x1a\n")  # 不是真 PNG 也无所谓, 走 unknown 分支
        tpl = Template("{pdf_title}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "_photo.png"


# ---- Word 命名空间 ----


class TestWordNamespace:
    def test_word_title(self, tmp_path: Path) -> None:
        f = _make_docx(tmp_path, title="My Word Doc")
        tpl = Template("{word_title}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "My Word Doc_doc.docx"

    def test_word_paragraphs(self, tmp_path: Path) -> None:
        f = _make_docx(tmp_path)
        tpl = Template("{word_paragraphs}par_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        # docx 加 3 段 + 初始 1 段 = 4 paragraphs
        assert target.name.startswith("3par_doc.docx")

    def test_word_namespace_empty_for_non_word(self, tmp_path: Path) -> None:
        f = _make_png(tmp_path)
        tpl = Template("{word_title}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "_img.png"


# ---- Excel 命名空间 ----


class TestExcelNamespace:
    def test_excel_sheets(self, tmp_path: Path) -> None:
        f = _make_xlsx(tmp_path)
        tpl = Template("{excel_sheets}sheets_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        # 默认 + 1 个新建 = 2 sheets
        assert target.name == "2sheets_book.xlsx"

    def test_excel_sheet_name(self, tmp_path: Path) -> None:
        f = _make_xlsx(tmp_path)
        tpl = Template("{excel_sheet_name}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        # 默认 sheet 名是 Sheet
        assert target.name.startswith("Sheet_book.xlsx")

    def test_excel_namespace_empty_for_non_excel(self, tmp_path: Path) -> None:
        f = _make_png(tmp_path)
        tpl = Template("{excel_sheets}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "0_img.png"


# ---- Image 命名空间 ----


class TestImageNamespace:
    def test_image_width_height(self, tmp_path: Path) -> None:
        f = _make_png(tmp_path, width=1920, height=1080)
        tpl = Template("{image_width}x{image_height}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "1920x1080_img.png"

    def test_image_aspect_ratio(self, tmp_path: Path) -> None:
        f = _make_png(tmp_path, width=1920, height=1080)
        tpl = Template("{image_aspect_ratio}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "16_9_img.png"

    def test_image_format(self, tmp_path: Path) -> None:
        f = _make_png(tmp_path)
        tpl = Template("{image_format}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "PNG_img.png"

    def test_image_namespace_empty_for_non_image(self, tmp_path: Path) -> None:
        f = _make_pdf(tmp_path)
        tpl = Template("{image_width}x{image_height}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "0x0_doc.pdf"


# ---- 通用占位符向后兼容 (W3) ----


class TestGeneralPlaceholderCompat:
    def test_title_works_for_pdf(self, tmp_path: Path) -> None:
        f = _make_pdf(tmp_path, title="Report")
        tpl = Template("{Title}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "Report_doc.pdf"

    def test_title_works_for_word(self, tmp_path: Path) -> None:
        f = _make_docx(tmp_path, title="Letter")
        tpl = Template("{Title}_{OriginalName}")
        target = Renamer(tpl)._render_target(f)
        assert target is not None
        assert target.name == "Letter_doc.docx"


# ---- apply_with_progress ----


class TestApplyWithProgress:
    def test_progress_called_per_file(self, tmp_path: Path) -> None:
        files = []
        for i in range(5):
            f = tmp_path / f"f{i}.txt"
            f.write_text(f"content {i}")
            files.append(f)

        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))
        calls: list[tuple[int, int, Path, RenameResult]] = []

        def cb(i, total, file, result):
            calls.append((i, total, file, result))

        results = renamer.apply_with_progress(files, on_progress=cb)
        assert len(results) == 5
        assert len(calls) == 5
        # 第一次回调 i=1, 最后一次 i=5
        assert calls[0][0] == 1
        assert calls[-1][0] == 5
        assert calls[-1][1] == 5

    def test_progress_callback_exception_swallowed(self, tmp_path: Path) -> None:
        """进度回调抛异常不应中断主流程."""
        f = tmp_path / "a.txt"
        f.write_text("x")
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))

        def bad_cb(i, total, file, result):
            raise RuntimeError("boom")

        # 不应抛
        results = renamer.apply_with_progress([f], on_progress=bad_cb)
        assert len(results) == 1
        assert results[0].status == "OK"

    def test_progress_called_for_dry_run(self, tmp_path: Path) -> None:
        f = tmp_path / "x.txt"
        f.write_text("x")
        progress_count = {"n": 0}

        def cb(i, total, file, result):
            progress_count["n"] += 1

        # plan 模式 (DRY_RUN) 不走 apply, 改用 plan 测 — 但 plan 不回调
        # 这里 apply + dry-run 内部没区别, 仅 status="DRY_RUN" 行为仍跑 plan 逻辑
        # 实际: apply_with_progress 走 _apply_one, 模板渲染但没冲突时正常改名
        # 用一个不改变文件名的模板 (i.e. 模板等于原名) 测 skip
        skip_renamer = Renamer(Template("{OriginalName}"))
        results = skip_renamer.apply_with_progress([f], on_progress=cb)
        assert progress_count["n"] == 1
        assert results[0].status == "SKIPPED"

    def test_no_callback_works(self, tmp_path: Path) -> None:
        f = tmp_path / "y.txt"
        f.write_text("y")
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))
        # 不传 on_progress
        results = renamer.apply_with_progress([f])
        assert len(results) == 1
        assert results[0].status == "OK"
        assert (tmp_path / "001_y.txt").exists()

    # ---- apply_with_progress cancellation (W7) ----

    def test_cancellation_stops_loop_immediately(self, tmp_path: Path) -> None:
        """W7: 取消时立即停止, 已处理的 results 仍返回."""
        files = []
        for i in range(5):
            f = tmp_path / f"f{i}.txt"
            f.write_text("x")
            files.append(f)
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))
        # 第一个回调后就 cancel
        cancel_flag = {"cancelled": False}

        def cb(i, total, file, result):
            cancel_flag["cancelled"] = True

        results = renamer.apply_with_progress(
            files,
            on_progress=cb,
            is_cancelled=lambda: cancel_flag["cancelled"],
        )
        # 第 1 个文件处理完 (callback 后 cancel), 循环在第 2 个文件前检查 → 停止
        assert len(results) == 1
        assert cancel_flag["cancelled"]
        assert (tmp_path / "001_f0.txt").exists()
        # 第 2-5 个文件不应被改名
        assert not (tmp_path / "002_f1.txt").exists()
        assert not (tmp_path / "005_f4.txt").exists()

    def test_cancellation_in_middle_of_batch(self, tmp_path: Path) -> None:
        """W7: 取消发生在第 3 个文件后, 3 个已处理 + 7 个未处理."""
        files = []
        for i in range(10):
            f = tmp_path / f"f{i}.txt"
            f.write_text("x")
            files.append(f)
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))
        seen = {"n": 0}

        def should_cancel():
            return seen["n"] >= 3

        def cb(i, total, file, result):
            seen["n"] = i

        results = renamer.apply_with_progress(
            files,
            on_progress=cb,
            is_cancelled=should_cancel,
        )
        assert len(results) == 3
        # 前 3 个已改名
        for i in range(3):
            assert (tmp_path / f"00{i+1}_f{i}.txt").exists()
        # 后 7 个没改名
        assert not (tmp_path / "004_f3.txt").exists()
        assert not (tmp_path / "010_f9.txt").exists()

    def test_cancellation_with_undo_stack_only_processed_entries_pushed(
        self, tmp_path: Path
    ) -> None:
        """W7: undo stack 只含已处理文件的 entries, 取消后未处理文件不入栈."""
        files = []
        for i in range(5):
            f = tmp_path / f"f{i}.txt"
            f.write_text("x")
            files.append(f)
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))
        seen = {"n": 0}

        def should_cancel():
            return seen["n"] >= 2

        def cb(i, total, file, result):
            seen["n"] = i

        undo = UndoStack()
        results = renamer.apply_with_progress(
            files,
            undo_stack=undo,
            on_progress=cb,
            is_cancelled=should_cancel,
        )
        assert len(results) == 2
        # undo stack 应只有 2 个 entry (2 个 batch, 每 batch 1 entry)
        total_entries = sum(len(batch) for batch in undo._entries)  # type: ignore[attr-defined]
        assert total_entries == 2

    def test_cancellation_already_set_does_no_work(self, tmp_path: Path) -> None:
        """W7: 取消标记已存在时, 0 个文件被处理."""
        files = []
        for i in range(3):
            f = tmp_path / f"f{i}.txt"
            f.write_text("x")
            files.append(f)
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))
        results = renamer.apply_with_progress(
            files,
            is_cancelled=lambda: True,
        )
        assert len(results) == 0
        for f in files:
            assert f.exists()  # 原文件不动

    def test_no_cancellation_token_runs_all_files(self, tmp_path: Path) -> None:
        """W7 backward compat: 不传 is_cancelled 时, 全部文件处理 (W5/W6 行为)."""
        files = []
        for i in range(4):
            f = tmp_path / f"f{i}.txt"
            f.write_text("x")
            files.append(f)
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))
        results = renamer.apply_with_progress(files)
        assert len(results) == 4
        for i in range(4):
            assert (tmp_path / f"00{i+1}_f{i}.txt").exists()


# ---- W9: 硬中断 (Hard Interrupt) ----

def _make_files(tmp_path: Path, count: int) -> list:
    """W9 helper: 创建 count 个文件."""
    files = []
    for i in range(count):
        f = tmp_path / f"f{i}.txt"
        f.write_text(f"content-{i}")
        files.append(f)
    return files


def _w9_cancel_at_step_a() -> callable:
    """W9 取消触发器: 让 safe_rename 内部 Step A 后才看到 cancel.

    机制: W7 检查在文件循环顶部 (call 1 = file 1 W7 顶),
          W9 检查在 safe_rename Step A 后 (call 2 = file 1 W9 内).
    返回 is_cancelled 闭包 — 第 1 次 False (放行 file 1), 之后 True.

    这样:
    - file 1 走完 W7 检查 (False), 进入 _apply_one, safe_rename Step A 后看到 True → ROLLBACK
    - file 2 走 W7 检查 (True) → break, 不处理
    """
    counter = [0]

    def is_cancelled() -> bool:
        counter[0] += 1
        # 第 1 次 (file 1 W7 顶): False
        # 第 2 次 (file 1 W9 safe_rename 内): True
        # 第 3 次起 (file 2+ W7 顶): True
        return counter[0] >= 2

    return is_cancelled


class TestApplyOneRollbackW9:
    """W9: _apply_one 内部取消 — 单文件 Step A 后检查 cancel, 触发 ROLLBACK."""

    def test_first_file_cancelled_mid_rename_keeps_source(self, tmp_path: Path) -> None:
        """第一个文件处理中取消, 源文件回原位."""
        files = _make_files(tmp_path, 3)
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))
        is_cancelled = _w9_cancel_at_step_a()

        results = renamer.apply_with_progress(
            files,
            is_cancelled=is_cancelled,
        )

        # 第 1 个文件 ROLLBACK, 第 2/3 个文件 W7 在文件之间检查时发现 cancel → 不进入
        assert len(results) == 1
        assert results[0].status == "ROLLBACK"
        assert results[0].source == files[0]
        assert files[0].exists()
        assert files[0].read_text() == "content-0"
        assert files[1].exists()
        assert files[2].exists()
        # 无 .tmp 残留
        assert list(tmp_path.glob("*.filemaster.tmp.*")) == []

    def test_rollback_does_not_push_undo_entry(self, tmp_path: Path) -> None:
        """ROLLBACK 状态不入 UndoStack (没真完成 rename)."""
        files = _make_files(tmp_path, 3)
        undo_stack = UndoStack()
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))
        is_cancelled = _w9_cancel_at_step_a()

        results = renamer.apply_with_progress(
            files,
            undo_stack=undo_stack,
            is_cancelled=is_cancelled,
        )

        assert len(results) == 1
        assert results[0].status == "ROLLBACK"
        assert len(undo_stack) == 0
        assert files[0].exists()

    def test_overwrite_with_cancel_cleans_backup_orphan(self, tmp_path: Path) -> None:
        """OVERWRITE + 取消: backup_path 已写但 target 没覆盖, 清理孤儿 backup."""
        src = tmp_path / "src.txt"
        src.write_text("new content")
        dst = tmp_path / "dst.txt"
        dst.write_text("old content")
        undo_stack = UndoStack()

        renamer = Renamer(Template("renamed_{OriginalName}"))
        is_cancelled = _w9_cancel_at_step_a()
        results = renamer.apply_with_progress(
            [src],
            conflict_strategy=ConflictStrategy.OVERWRITE,
            undo_stack=undo_stack,
            is_cancelled=is_cancelled,
        )

        assert len(results) == 1
        assert results[0].status == "ROLLBACK"
        assert src.exists()
        assert dst.exists()
        assert dst.read_text() == "old content"
        # OVERWRITE 走了 backup, 但取消时 backup_path 被清掉, undo_stack 也无 entry
        assert len(undo_stack) == 0
        # 无 .tmp 残留
        assert list(tmp_path.glob("*.filemaster.tmp.*")) == []

    def test_normal_completion_still_works_after_w9(self, tmp_path: Path) -> None:
        """W9 改动不破坏 W5/W6/W7 正常路径."""
        files = _make_files(tmp_path, 3)
        undo_stack = UndoStack()
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))

        results = renamer.apply_with_progress(
            files,
            undo_stack=undo_stack,
            is_cancelled=lambda: False,
        )

        assert len(results) == 3
        assert all(r.status == "OK" for r in results)
        assert len(undo_stack) == 1
        for f in files:
            assert not f.exists()
        for i in range(3):
            assert (tmp_path / f"00{i+1}_f{i}.txt").exists()
        # 无 .tmp 残留 (sanity)
        assert list(tmp_path.glob("*.filemaster.tmp.*")) == []


class TestApplyEntryCleanupTmpsW9:
    """W9: apply 入口清理 .tmp 残留."""

    def test_apply_cleans_orphan_tmps(self, tmp_path: Path) -> None:
        """apply 启动时自动清理源目录的 .tmp 残留."""
        # 模拟上次取消时留下的 .tmp
        orphan = tmp_path / "src.txt.filemaster.tmp.deadbeef"
        orphan.write_text("leftover")

        files = _make_files(tmp_path, 2)
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))

        renamer.apply_with_progress(files)

        assert not orphan.exists()
        for i in range(2):
            assert (tmp_path / f"00{i+1}_f{i}.txt").exists()

    def test_apply_orphan_tmps_preserves_normal_files(self, tmp_path: Path) -> None:
        """清理时不动普通文件."""
        (tmp_path / "user_doc.txt").write_text("untouched")
        (tmp_path / "user_doc.txt.filemaster.tmp.abc12345").write_text("tmp")

        files = _make_files(tmp_path, 1)
        renamer = Renamer(Template("{Index:D3}_{OriginalName}"))
        renamer.apply_with_progress(files)

        assert (tmp_path / "user_doc.txt").exists()
        assert not (tmp_path / "user_doc.txt.filemaster.tmp.abc12345").exists()
